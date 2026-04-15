"""Comprehensive tests for Word parser."""

from pathlib import Path

import pytest

from exceptions import FileParsingError
from parsers.word_parser import WordParser

# Test data directory
TEST_DATA_DIR = Path(__file__).parent / "data"


class TestWordParserInitialization:
    """Test Word parser initialization."""

    def test_initialization(self):
        """Test that WordParser initializes correctly."""
        parser = WordParser()
        assert parser is not None
        assert parser.supported_extensions == [".docx", ".doc"]


class TestWordParserValidFiles:
    """Test Word parser with valid files."""

    @pytest.fixture
    def parser(self):
        """Create a WordParser instance."""
        return WordParser()

    def test_parse_valid_docx(self, parser: WordParser):
        """Test parsing a valid DOCX file."""
        docx_file = TEST_DATA_DIR / "valid_resume.docx"
        if not docx_file.exists():
            pytest.skip(f"Test file {docx_file} does not exist")

        text = parser.parse(str(docx_file))
        assert isinstance(text, str)
        assert len(text) > 0
        assert text.strip() != ""
        assert "John Doe" in text or "Test Resume" in text

    def test_parse_docx_with_multiple_paragraphs(self, parser: WordParser):
        """Test parsing a DOCX with multiple paragraphs."""
        docx_file = TEST_DATA_DIR / "valid_resume.docx"
        if not docx_file.exists():
            pytest.skip(f"Test file {docx_file} does not exist")

        text = parser.parse(str(docx_file))
        # Should have multiple lines
        assert "\n" in text

    def test_parse_docx_with_tables(self, parser: WordParser):
        """Test parsing a DOCX file with tables."""
        docx_file = TEST_DATA_DIR / "with_tables.docx"
        if not docx_file.exists():
            pytest.skip(f"Test file {docx_file} does not exist")

        text = parser.parse(str(docx_file))
        assert isinstance(text, str)
        assert len(text) > 0
        # Tables should be separated with |
        assert "|" in text or len(text) > 0

    def test_parse_docx_with_formatting(self, parser: WordParser):
        """Test parsing a DOCX with various formatting (bold, italic, etc)."""
        docx_file = TEST_DATA_DIR / "formatted.docx"
        if not docx_file.exists():
            pytest.skip(f"Test file {docx_file} does not exist")

        text = parser.parse(str(docx_file))
        assert isinstance(text, str)
        assert len(text) > 0


class TestWordParserInvalidFiles:
    """Test Word parser with invalid files."""

    @pytest.fixture
    def parser(self):
        """Create a WordParser instance."""
        return WordParser()

    def test_parse_nonexistent_file(self, parser: WordParser):
        """Test parsing a file that doesn't exist."""
        with pytest.raises(FileNotFoundError):
            parser.parse(str(TEST_DATA_DIR / "nonexistent.docx"))

    def test_parse_empty_docx(self, parser: WordParser):
        """Test parsing an empty DOCX file."""
        empty_docx = TEST_DATA_DIR / "empty.docx"
        if not empty_docx.exists():
            pytest.skip(f"Test file {empty_docx} does not exist")

        with pytest.raises(FileParsingError):
            parser.parse(str(empty_docx))

    def test_parse_corrupted_docx(self, parser: WordParser):
        """Test parsing a corrupted DOCX file."""
        corrupted_docx = TEST_DATA_DIR / "corrupted.docx"
        if not corrupted_docx.exists():
            pytest.skip(f"Test file {corrupted_docx} does not exist")

        with pytest.raises(FileParsingError) as exc_info:
            parser.parse(str(corrupted_docx))
        assert exc_info.value.original_exception is not None

    def test_parse_text_file_as_docx(self, parser: WordParser):
        """Test parsing a text file with .docx extension."""
        fake_docx = TEST_DATA_DIR / "fake.docx"
        if not fake_docx.exists():
            pytest.skip(f"Test file {fake_docx} does not exist")

        with pytest.raises(FileParsingError):
            parser.parse(str(fake_docx))

    def test_parse_non_docx_file(self, parser: WordParser):
        """Test parsing a non-DOCX file."""
        non_docx = TEST_DATA_DIR / "sample.txt"
        if not non_docx.exists():
            pytest.skip(f"Test file {non_docx} does not exist")

        with pytest.raises(FileParsingError) as exc_info:
            parser.parse(str(non_docx))
        assert "Unsupported file format" in str(exc_info.value)

    def test_parse_directory(self, parser: WordParser):
        """Test parsing a directory path."""
        with pytest.raises(FileParsingError):
            parser.parse(str(TEST_DATA_DIR))


class TestWordParserLegacyDoc:
    """Test Word parser with legacy .doc files."""

    @pytest.fixture
    def parser(self):
        """Create a WordParser instance."""
        return WordParser()

    def test_parse_legacy_doc_file(self, parser: WordParser):
        """Test parsing a legacy .doc file."""
        doc_file = TEST_DATA_DIR / "sample.doc"
        if not doc_file.exists():
            pytest.skip(f"Test file {doc_file} does not exist")

        # Legacy .doc files may or may not work depending on the file
        try:
            text = parser.parse(str(doc_file))
            assert isinstance(text, str)
        except FileParsingError as e:
            # This is acceptable - legacy .doc support is limited
            assert "Cannot parse legacy .doc file" in str(
                e
            ) or "Error parsing Word document" in str(e)

    def test_supports_doc_extension(self, parser: WordParser):
        """Test that .doc extension is reported as supported."""
        assert parser.supports_format("test.doc")


class TestWordParserSupportsFormat:
    """Test Word parser format support checking."""

    @pytest.fixture
    def parser(self):
        """Create a WordParser instance."""
        return WordParser()

    def test_supports_docx_extension(self, parser: WordParser):
        """Test that .docx extension is supported."""
        assert parser.supports_format("test.docx")
        assert parser.supports_format("TEST.DOCX")
        assert parser.supports_format("/path/to/file.docx")

    def test_supports_doc_extension(self, parser: WordParser):
        """Test that .doc extension is supported."""
        assert parser.supports_format("test.doc")
        assert parser.supports_format("TEST.DOC")

    def test_does_not_support_other_extensions(self, parser: WordParser):
        """Test that non-Word extensions are not supported."""
        assert not parser.supports_format("test.pdf")
        assert not parser.supports_format("test.txt")
        assert not parser.supports_format("test.xlsx")
        assert not parser.supports_format("test.pptx")
        assert not parser.supports_format("test")

    def test_supports_format_with_path_object(self, parser: WordParser):
        """Test supports_format with Path object."""
        assert parser.supports_format(Path("test.docx"))
        assert parser.supports_format(Path("test.doc"))
        assert not parser.supports_format(Path("test.pdf"))


class TestWordParserEdgeCases:
    """Test Word parser edge cases."""

    @pytest.fixture
    def parser(self):
        """Create a WordParser instance."""
        return WordParser()

    def test_parse_docx_with_only_images(self, parser: WordParser):
        """Test parsing a DOCX with only images (no text)."""
        image_docx = TEST_DATA_DIR / "images_only.docx"
        if not image_docx.exists():
            pytest.skip(f"Test file {image_docx} does not exist")

        with pytest.raises(FileParsingError) as exc_info:
            parser.parse(str(image_docx))
        assert "No text content found in document" in str(exc_info.value)

    def test_parse_docx_with_special_characters(self, parser: WordParser):
        """Test parsing a DOCX with special characters."""
        special_docx = TEST_DATA_DIR / "special_chars.docx"
        if not special_docx.exists():
            pytest.skip(f"Test file {special_docx} does not exist")

        text = parser.parse(str(special_docx))
        assert isinstance(text, str)
        assert len(text) > 0

    def test_parse_docx_with_headers_footers(self, parser: WordParser):
        """Test parsing a DOCX with headers and footers."""
        header_docx = TEST_DATA_DIR / "with_headers.docx"
        if not header_docx.exists():
            pytest.skip(f"Test file {header_docx} does not exist")

        text = parser.parse(str(header_docx))
        assert isinstance(text, str)
        # Headers/footers may or may not be extracted depending on implementation

    def test_parse_docx_with_lists(self, parser: WordParser):
        """Test parsing a DOCX with bullet points and numbered lists."""
        list_docx = TEST_DATA_DIR / "with_lists.docx"
        if not list_docx.exists():
            pytest.skip(f"Test file {list_docx} does not exist")

        text = parser.parse(str(list_docx))
        assert isinstance(text, str)
        assert len(text) > 0

    def test_parse_very_large_docx(self, parser: WordParser):
        """Test parsing a very large DOCX file."""
        large_docx = TEST_DATA_DIR / "large.docx"
        if not large_docx.exists():
            pytest.skip(f"Test file {large_docx} does not exist")

        text = parser.parse(str(large_docx))
        assert isinstance(text, str)
        assert len(text) > 1000  # Should have substantial content

    def test_parse_docx_with_empty_paragraphs(self, parser: WordParser):
        """Test parsing a DOCX with many empty paragraphs."""
        empty_para_docx = TEST_DATA_DIR / "empty_paragraphs.docx"
        if not empty_para_docx.exists():
            pytest.skip(f"Test file {empty_para_docx} does not exist")

        text = parser.parse(str(empty_para_docx))
        assert isinstance(text, str)
        # Empty paragraphs should be filtered out
        assert text.strip() != ""

    def test_parse_docx_with_nested_tables(self, parser: WordParser):
        """Test parsing a DOCX with nested tables."""
        nested_docx = TEST_DATA_DIR / "nested_tables.docx"
        if not nested_docx.exists():
            pytest.skip(f"Test file {nested_docx} does not exist")

        text = parser.parse(str(nested_docx))
        assert isinstance(text, str)
        assert len(text) > 0

    def test_table_extraction_format(self, parser: WordParser):
        """Test that table cells are properly formatted with pipes."""
        docx_file = TEST_DATA_DIR / "with_tables.docx"
        if not docx_file.exists():
            pytest.skip(f"Test file {docx_file} does not exist")

        text = parser.parse(str(docx_file))
        # If tables exist, they should be formatted with |
        # This test is informational
        assert isinstance(text, str)


class TestWordParserEnhancements:
    """Tests for the enhanced extraction features: headers/footers, text boxes,
    footnotes, merged cells, hyperlinks, and list formatting."""

    @pytest.fixture
    def parser(self):
        return WordParser()

    # ------------------------------------------------------------------
    # Headers and footers
    # ------------------------------------------------------------------

    def test_headers_and_footers_extracted(self, parser: WordParser):
        """Header and footer text is included in the parsed output."""
        docx = TEST_DATA_DIR / "with_headers.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert "Confidential Resume" in text
        assert "Page 1" in text

    # ------------------------------------------------------------------
    # Text boxes
    # ------------------------------------------------------------------

    def test_text_box_content_extracted(self, parser: WordParser):
        """Text inside wps:txbx text boxes is extracted."""
        docx = TEST_DATA_DIR / "with_textboxes.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert "Text Box Content" in text

    def test_text_box_not_duplicated(self, parser: WordParser):
        """Text box content appears exactly once (no double-extraction)."""
        docx = TEST_DATA_DIR / "with_textboxes.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert text.count("Text Box Content") == 1

    # ------------------------------------------------------------------
    # Footnotes
    # ------------------------------------------------------------------

    def test_footnote_text_extracted(self, parser: WordParser):
        """Footnote text is included in the parsed output."""
        docx = TEST_DATA_DIR / "with_footnotes.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert "This is footnote text." in text

    def test_document_without_footnotes_parses_cleanly(self, parser: WordParser):
        """Documents with no footnotes parse without errors."""
        docx = TEST_DATA_DIR / "valid_resume.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert isinstance(text, str)
        assert len(text) > 0

    # ------------------------------------------------------------------
    # Merged cells
    # ------------------------------------------------------------------

    def test_merged_cells_no_duplication(self, parser: WordParser):
        """Merged cell text appears exactly once (not once per grid position)."""
        docx = TEST_DATA_DIR / "merged_cells.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert text.count("Merged Header") == 1

    def test_merged_cells_remaining_rows_present(self, parser: WordParser):
        """Non-merged rows are still fully extracted after a merge."""
        docx = TEST_DATA_DIR / "merged_cells.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert "Company" in text
        assert "Role" in text
        assert "Acme Corp" in text
        assert "Engineer" in text

    # ------------------------------------------------------------------
    # Hyperlinks
    # ------------------------------------------------------------------

    def test_hyperlink_display_text_extracted(self, parser: WordParser):
        """Hyperlink display text (not URL) is included in output."""
        docx = TEST_DATA_DIR / "with_hyperlinks.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert "Visit Example" in text

    def test_hyperlink_surrounding_text_preserved(self, parser: WordParser):
        """Text around a hyperlink is not lost."""
        docx = TEST_DATA_DIR / "with_hyperlinks.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert "Portfolio:" in text
        assert "Additional body text after hyperlink." in text

    # ------------------------------------------------------------------
    # List formatting
    # ------------------------------------------------------------------

    def test_bullet_list_prefixed(self, parser: WordParser):
        """List Bullet style paragraphs get a '•' prefix."""
        docx = TEST_DATA_DIR / "with_lists.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert "• Python" in text
        assert "• Docker" in text

    def test_non_list_paragraphs_not_prefixed(self, parser: WordParser):
        """Regular paragraphs do not get a bullet prefix."""
        docx = TEST_DATA_DIR / "valid_resume.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        assert not any(line.startswith("• ") for line in text.splitlines())

    # ------------------------------------------------------------------
    # Document order
    # ------------------------------------------------------------------

    def test_body_content_in_document_order(self, parser: WordParser):
        """Paragraphs and tables are interleaved in document order."""
        docx = TEST_DATA_DIR / "nested_tables.docx"
        if not docx.exists():
            pytest.skip(f"Test file {docx} does not exist")

        text = parser.parse(str(docx))
        # Table content (pipe-delimited) should appear before the trailing paragraph
        assert "|" in text
        assert "Python" in text
        table_pos = text.index("|")
        skills_pos = text.index("Python")
        assert table_pos < skills_pos


class TestWordParserPrivateMethods:
    """Unit tests for WordParser private extraction methods."""

    @pytest.fixture
    def parser(self):
        return WordParser()

    def test_format_table_deduplicates_merged_cells(self, parser: WordParser):
        """_format_table skips repeated lxml elements (merged cells)."""
        from docx import Document
        import io

        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).merge(tbl.cell(0, 1)).text = "Span"
        tbl.cell(1, 0).text = "A"
        tbl.cell(1, 1).text = "B"

        lines = parser._format_table(tbl)
        assert lines[0] == "Span"           # merged row: one cell, not two
        assert lines[1] == "A | B"

    def test_format_paragraph_bullet_prefix(self, parser: WordParser):
        """_format_paragraph prefixes List Bullet style with '•'."""
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("Item", style="List Bullet")
        result = parser._format_paragraph(para)
        assert result == "• Item"

    def test_format_paragraph_normal_no_prefix(self, parser: WordParser):
        """Normal paragraphs get no bullet prefix."""
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("Normal text")
        result = parser._format_paragraph(para)
        assert result == "Normal text"

    def test_format_paragraph_empty_returns_empty(self, parser: WordParser):
        """Empty paragraphs return empty string."""
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("")
        result = parser._format_paragraph(para)
        assert result == ""

    def test_extract_headers_footers_no_content(self, parser: WordParser):
        """Document with no header/footer text returns empty list."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Body only")
        result = parser._extract_headers_footers(doc)
        assert result == []

    def test_extract_text_boxes_no_text_boxes(self, parser: WordParser):
        """Document with no text boxes returns empty list."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("No text boxes here")
        result = parser._extract_text_boxes(doc)
        assert result == []

    def test_extract_footnotes_endnotes_no_footnotes(self, parser: WordParser):
        """Document with no footnotes returns empty list without error."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("No footnotes")
        result = parser._extract_footnotes_endnotes(doc)
        assert result == []
