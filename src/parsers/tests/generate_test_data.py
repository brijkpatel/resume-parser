"""Script to generate test data files for parser tests."""

from pathlib import Path
from docx import Document

# Create test data directory
data_dir = Path(__file__).parent / "data"
data_dir.mkdir(exist_ok=True)


# Create a simple valid DOCX file
def create_simple_docx():
    doc = Document()
    doc.add_heading("Test Resume", 0)
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Email: john.doe@example.com")
    doc.add_paragraph("Phone: +1-234-567-8900")

    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Software Engineer at Tech Company")
    doc.add_paragraph("Worked on various projects using Python, Java, and JavaScript.")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Java, JavaScript, SQL, Git")

    # Save the document
    doc.save(str(data_dir / "valid_resume.docx"))
    print(f"Created valid_resume.docx")


# Create a DOCX with tables
def create_docx_with_tables():
    doc = Document()
    doc.add_heading("Resume with Tables", 0)

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"

    # Fill the table
    cells = [
        ("Name", "Jane Smith"),
        ("Email", "jane@example.com"),
        ("Phone", "+1-987-654-3210"),
    ]

    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Data Analysis, Machine Learning")

    doc.save(str(data_dir / "with_tables.docx"))
    print(f"Created with_tables.docx")


# Create a DOCX with only whitespace/empty paragraphs
def create_empty_content_docx():
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("\n")
    doc.save(str(data_dir / "empty_content.docx"))
    print(f"Created empty_content.docx")


def create_formatted_docx():
    """Create a DOCX with bold, italic, and underlined text."""
    doc = Document()
    doc.add_heading("Formatted Resume", 0)

    p = doc.add_paragraph()
    run = p.add_run("John Bold")
    run.bold = True

    p2 = doc.add_paragraph()
    run2 = p2.add_run("Senior Software Engineer")
    run2.italic = True

    p3 = doc.add_paragraph()
    run3 = p3.add_run("Contact: john@example.com")
    run3.underline = True

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Java, Go, Kubernetes")

    doc.save(str(data_dir / "formatted.docx"))
    print("Created formatted.docx")


def create_images_only_docx():
    """Create a DOCX whose paragraphs contain only whitespace.

    The parser strips every paragraph; finding no text it raises
    FileParsingError("No text content found in document.").
    """
    doc = Document()
    for _ in range(5):
        doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.save(str(data_dir / "images_only.docx"))
    print("Created images_only.docx")


def create_special_chars_docx():
    """Create a DOCX with international and special characters."""
    doc = Document()
    doc.add_heading("Résumé — Special Characters", 0)
    doc.add_paragraph("Name: José García-Müller")
    doc.add_paragraph("Email: josé.garcía@example.com")
    doc.add_paragraph("Skills: C++, C#, Python, Rust")
    doc.add_paragraph("Languages: English, Español, Français, 中文")
    doc.add_paragraph("Note: Salary expectation €75,000–€90,000")

    doc.save(str(data_dir / "special_chars.docx"))
    print("Created special_chars.docx")


def create_headers_docx():
    """Create a DOCX with a document header and footer."""
    doc = Document()

    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Confidential Resume — Jane Smith"

    footer = section.footer
    footer.paragraphs[0].text = "Page 1"

    doc.add_heading("Jane Smith", 0)
    doc.add_paragraph("Software Engineer | jane@example.com")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Lead Engineer at Acme Corp, 2020–present")

    doc.save(str(data_dir / "with_headers.docx"))
    print("Created with_headers.docx")


def create_lists_docx():
    """Create a DOCX with bullet-point and numbered-list paragraphs."""
    doc = Document()
    doc.add_heading("Resume with Lists", 0)

    doc.add_heading("Skills", level=1)
    for skill in ["Python", "Java", "Docker", "Kubernetes", "PostgreSQL"]:
        doc.add_paragraph(skill, style="List Bullet")

    doc.add_heading("Achievements", level=1)
    for i, achievement in enumerate(
        ["Reduced latency by 40%", "Led team of 8 engineers", "Shipped 3 major releases"],
        start=1,
    ):
        doc.add_paragraph(achievement, style="List Number")

    doc.save(str(data_dir / "with_lists.docx"))
    print("Created with_lists.docx")


def create_large_docx():
    """Create a DOCX with substantial text content (>1000 chars)."""
    doc = Document()
    doc.add_heading("Large Resume Document", 0)

    for i in range(1, 51):
        doc.add_paragraph(
            f"Section {i}: Experienced engineer with deep expertise in distributed "
            f"systems, cloud infrastructure, and high-performance data pipelines. "
            f"Delivered projects serving millions of users across multiple regions."
        )

    doc.save(str(data_dir / "large.docx"))
    print("Created large.docx")


def create_empty_paragraphs_docx():
    """Create a DOCX with many empty paragraphs interspersed with real content."""
    doc = Document()
    doc.add_heading("Resume", 0)
    doc.add_paragraph("")  # empty
    doc.add_paragraph("   ")  # whitespace only
    doc.add_paragraph("John Doe — Software Engineer")
    doc.add_paragraph("")
    doc.add_paragraph("john.doe@example.com | +1-555-000-1234")
    doc.add_paragraph("   ")
    doc.add_paragraph("")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("")
    doc.add_paragraph("Python, SQL, Docker")

    doc.save(str(data_dir / "empty_paragraphs.docx"))
    print("Created empty_paragraphs.docx")


def create_nested_tables_docx():
    """Create a DOCX with a table containing multiple rows and columns."""
    doc = Document()
    doc.add_heading("Resume with Nested Tables", 0)

    # Outer table representing work history
    table = doc.add_table(rows=3, cols=3)
    headers = ["Company", "Role", "Duration"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h

    rows_data = [
        ("Acme Corp", "Senior Engineer", "2020–present"),
        ("Beta Inc", "Software Engineer", "2017–2020"),
    ]
    for i, (company, role, duration) in enumerate(rows_data, start=1):
        table.rows[i].cells[0].text = company
        table.rows[i].cells[1].text = role
        table.rows[i].cells[2].text = duration

    doc.add_paragraph()
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Java, Kubernetes")

    doc.save(str(data_dir / "nested_tables.docx"))
    print("Created nested_tables.docx")


def create_legacy_doc():
    """Create a stub .doc file that python-docx cannot parse.

    python-docx requires a ZIP-format container (.docx); a .doc file in
    legacy binary format causes Document() to raise BadZipFile, which the
    word parser wraps as FileParsingError("Error parsing Word document.").
    """
    # Write an OLE2 compound-document magic header followed by zeroed bytes.
    # This is recognisable as a legacy Word binary file but python-docx will
    # reject it, satisfying the test's error-handling assertion.
    ole2_magic = b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"
    (data_dir / "sample.doc").write_bytes(ole2_magic + b"\x00" * 504)
    print("Created sample.doc")


def create_textboxes_docx():
    """Create a DOCX with a modern wps:txbx text box containing 'Text Box Content'.

    The text box is injected as raw XML since python-docx has no public API
    for creating text boxes.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from lxml import etree

    doc = Document()
    doc.add_heading("Resume with Text Box", 0)
    doc.add_paragraph("Normal body paragraph.")

    # Build the wps:txbx XML fragment.
    txbx_xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
        ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
        "<w:r><w:rPr/></w:r>"
        "<w:r>"
        "<w:drawing>"
        '<wp:inline distT="0" distB="0" distL="0" distR="0">'
        '<wp:extent cx="2000000" cy="1000000"/>'
        "<wp:docPr id=\"1\" name=\"TextBox1\"/>"
        "<a:graphic>"
        "<a:graphicData"
        ' uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<wps:wsp>"
        "<wps:txbx>"
        '<w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:p><w:r><w:t>Text Box Content</w:t></w:r></w:p>"
        "</w:txbxContent>"
        "</wps:txbx>"
        "</wps:wsp>"
        "</a:graphicData>"
        "</a:graphic>"
        "</wp:inline>"
        "</w:drawing>"
        "</w:r>"
        "</w:p>"
    )
    txbx_elem = etree.fromstring(txbx_xml)
    doc.element.body.append(txbx_elem)

    doc.save(str(data_dir / "with_textboxes.docx"))
    print("Created with_textboxes.docx")


def create_merged_cells_docx():
    """Create a DOCX with a table that has merged cells.

    Row 0: cells (0,0) and (0,1) are merged → single cell text "Merged Header".
    Row 1: "Company" | "Role"
    Row 2: "Acme Corp" | "Engineer"
    Tests assert "Merged Header" appears exactly once (no duplication).
    """
    doc = Document()
    doc.add_heading("Resume — Merged Table", 0)

    table = doc.add_table(rows=3, cols=2)

    # Merge first row across both columns.
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Merged Header"

    table.cell(1, 0).text = "Company"
    table.cell(1, 1).text = "Role"
    table.cell(2, 0).text = "Acme Corp"
    table.cell(2, 1).text = "Engineer"

    doc.save(str(data_dir / "merged_cells.docx"))
    print("Created merged_cells.docx")


def create_hyperlinks_docx():
    """Create a DOCX paragraph containing a hyperlink with display text.

    The hyperlink display text is "Visit Example"; the URL target is
    https://example.com.  Tests assert "Visit Example" appears in parsed output.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document()
    doc.add_heading("Resume with Hyperlinks", 0)
    doc.add_paragraph("Contact information below.")

    # Add a relationship for the hyperlink.
    para = doc.add_paragraph("Portfolio: ")
    part = doc.part
    r_id = part.relate_to(
        "https://example.com",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    # Build <w:hyperlink r:id="..."><w:r><w:t>Visit Example</w:t></w:r></w:hyperlink>
    hyperlink_xml = (
        f'<w:hyperlink r:id="{r_id}"'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        "<w:r>"
        "<w:rPr><w:rStyle w:val=\"Hyperlink\"/></w:rPr>"
        "<w:t>Visit Example</w:t>"
        "</w:r>"
        "</w:hyperlink>"
    )
    hl_elem = etree.fromstring(hyperlink_xml)
    para._element.append(hl_elem)

    doc.add_paragraph("Additional body text after hyperlink.")
    doc.save(str(data_dir / "with_hyperlinks.docx"))
    print("Created with_hyperlinks.docx")


def create_footnotes_docx():
    """Create a DOCX with a footnote by directly manipulating the XML parts.

    The footnote text is "This is footnote text."
    Tests assert this text appears in the parsed output.
    """
    import zipfile
    import shutil
    import io
    import re

    # Build a simple docx from python-docx first.
    doc = Document()
    doc.add_heading("Resume with Footnote", 0)
    doc.add_paragraph("Main body text.")

    # Save to an in-memory buffer so we can rewrite the zip.
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    footnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:footnotes xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">\n'
        '  <w:footnote w:type="separator" w:id="-1">'
        "<w:p><w:r><w:separator/></w:r></w:p></w:footnote>\n"
        '  <w:footnote w:type="continuationSeparator" w:id="0">'
        "<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>\n"
        '  <w:footnote w:id="1">'
        "<w:p><w:r><w:t>This is footnote text.</w:t></w:r></w:p>"
        "</w:footnote>\n"
        "</w:footnotes>"
    )

    # Rewrite the zip, injecting the footnotes part and updating [Content_Types].xml
    # and word/_rels/document.xml.rels.
    out_buf = io.BytesIO()
    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                content = data.decode("utf-8")
                if "footnotes" not in content:
                    insert = (
                        '<Override PartName="/word/footnotes.xml"'
                        ' ContentType="application/vnd.openxmlformats-officedocument'
                        '.wordprocessingml.footnotes+xml"/>'
                    )
                    content = content.replace("</Types>", insert + "</Types>")
                data = content.encode("utf-8")

            elif item.filename == "word/_rels/document.xml.rels":
                content = data.decode("utf-8")
                if "footnotes" not in content:
                    rel = (
                        '<Relationship Id="rFootnotes" Type="http://schemas.openxmlformats.org'
                        '/officeDocument/2006/relationships/footnotes"'
                        ' Target="footnotes.xml"/>'
                    )
                    content = content.replace("</Relationships>", rel + "</Relationships>")
                data = content.encode("utf-8")

            zout.writestr(item, data)

        # Add the new footnotes.xml part.
        zout.writestr("word/footnotes.xml", footnotes_xml.encode("utf-8"))

    out_path = data_dir / "with_footnotes.docx"
    out_path.write_bytes(out_buf.getvalue())
    print("Created with_footnotes.docx")


if __name__ == "__main__":
    try:
        create_simple_docx()
        create_docx_with_tables()
        create_empty_content_docx()
        create_formatted_docx()
        create_images_only_docx()
        create_special_chars_docx()
        create_headers_docx()
        create_lists_docx()
        create_large_docx()
        create_empty_paragraphs_docx()
        create_nested_tables_docx()
        create_legacy_doc()
        create_textboxes_docx()
        create_merged_cells_docx()
        create_hyperlinks_docx()
        create_footnotes_docx()
        print("\nAll test DOCX files created successfully!")
    except Exception as e:
        print(f"Error creating test files: {e}")
