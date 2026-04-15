"""Extract text from Word documents (.docx and .doc files)."""

from pathlib import Path
from typing import List, Union

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from exceptions import FileParsingError
from interfaces import FileParser
from utils import logger

# Fully-qualified XML namespaces used when iterating raw lxml elements.
_WPS_TXBX = (
    "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"
)
_VML_TEXTBOX = "{urn:schemas-microsoft-com:vml}textbox"
_W_T = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
_W_P = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
_W_FOOTNOTE = (
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnote"
)
_W_ENDNOTE = (
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}endnote"
)
_W_ID = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id"


class WordParser(FileParser):
    """Parse Word documents and extract text content.

    Extracts from (in order):
    1. Headers and footers (all sections).
    2. Body content — paragraphs and tables in document order.
    3. Text boxes (modern wps:txbx and legacy VML v:textbox).
    4. Footnotes and endnotes.

    Improvements over the basic version:
    - Document order preserved (paragraphs and tables interleaved correctly).
    - Merged table cells deduplicated.
    - Hyperlink display text captured.
    - List bullets prefixed with "•".
    - Text boxes and footnotes included.
    """

    def __init__(self):
        self.supported_extensions = [".docx", ".doc"]
        logger.debug("WordParser initialized")

    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------

    def parse(self, file_path: str) -> str:
        """Extract text from a Word document.

        Args:
            file_path: Path to the .docx or .doc file.

        Returns:
            Extracted text content.

        Raises:
            FileParsingError: If the file cannot be parsed.
        """
        self._validate_file_path(file_path)

        if not self.supports_format(file_path):
            raise FileParsingError("Unsupported file format.")

        logger.info("Parsing Word document: %s", file_path)

        if self._get_file_extension(file_path) == ".doc":
            return self._parse_doc_file(file_path)

        try:
            document = Document(file_path)
            text_parts: List[str] = []

            text_parts.extend(self._extract_headers_footers(document))
            text_parts.extend(self._extract_body_content(document))
            text_parts.extend(self._extract_text_boxes(document))
            text_parts.extend(self._extract_footnotes_endnotes(document))

            if not text_parts:
                raise FileParsingError("No text content found in document.")

            full_text = "\n".join(text_parts)
            logger.info(
                "Successfully parsed Word document: %s (%d characters)",
                file_path,
                len(full_text),
            )
            return full_text

        except FileParsingError:
            raise
        except Exception as e:
            raise FileParsingError(
                "Error parsing Word document.", original_exception=e
            ) from e

    def supports_format(self, file_path: Union[str, Path]) -> bool:
        """Return True if the file is a Word document."""
        return self._get_file_extension(file_path) in self.supported_extensions

    # ------------------------------------------------------------------
    # Headers and footers
    # ------------------------------------------------------------------

    def _extract_headers_footers(self, document: Document) -> List[str]:
        """Extract text from all section headers and footers."""
        texts: List[str] = []
        for section in document.sections:
            for hf in (
                section.header,
                section.footer,
                section.even_page_header,
                section.even_page_footer,
                section.first_page_header,
                section.first_page_footer,
            ):
                if hf.is_linked_to_previous:
                    continue
                for para in hf.paragraphs:
                    t = para.text.strip()
                    if t:
                        texts.append(t)
        return texts

    # ------------------------------------------------------------------
    # Body content (document order)
    # ------------------------------------------------------------------

    def _extract_body_content(self, document: Document) -> List[str]:
        """Walk body elements in document order, preserving paragraph/table interleave."""
        texts: List[str] = []
        for child in document.element.body:
            local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if local == "p":
                para = Paragraph(child, document)
                t = self._format_paragraph(para)
                if t:
                    texts.append(t)
            elif local == "tbl":
                tbl = Table(child, document)
                texts.extend(self._format_table(tbl))
        return texts

    def _format_paragraph(self, para: Paragraph) -> str:
        """Return paragraph text, including hyperlink runs and list bullets.

        Runs that contain drawings (images, text boxes) are skipped — their
        text is handled separately by _extract_text_boxes to avoid duplication.
        """
        parts: List[str] = []
        for child in para._element:
            local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if local == "r":
                # Skip drawing containers — text boxes are extracted separately.
                if any("drawing" in c.tag or "pict" in c.tag for c in child):
                    continue
                t = "".join(node.text or "" for node in child.iter(qn("w:t")))
                if t:
                    parts.append(t)
            elif local == "hyperlink":
                for run in child.iter(qn("w:r")):
                    t = "".join(node.text or "" for node in run.iter(qn("w:t")))
                    if t:
                        parts.append(t)

        text = "".join(parts).strip()
        if not text:
            return ""

        style_name = para.style.name if para.style else ""
        if "List Bullet" in style_name:
            return f"• {text}"
        return text

    def _format_table(self, table: Table) -> List[str]:
        """Format table rows as pipe-delimited strings, deduplicating merged cells.

        python-docx exposes the same lxml element for every position a merged
        cell occupies, so we track element identity (not id()) to skip dupes.
        """
        lines: List[str] = []
        seen: set = set()
        for row in table.rows:
            row_texts: List[str] = []
            for cell in row.cells:
                elem = cell._element
                if elem in seen:
                    continue
                seen.add(elem)
                t = cell.text.strip()
                if t:
                    row_texts.append(t)
            if row_texts:
                lines.append(" | ".join(row_texts))
        return lines

    # ------------------------------------------------------------------
    # Text boxes
    # ------------------------------------------------------------------

    def _extract_text_boxes(self, document: Document) -> List[str]:
        """Extract text from modern (wps:txbx) and legacy (v:textbox) text boxes."""
        texts: List[str] = []
        body = document.element.body
        for txbx in body.iter(_WPS_TXBX, _VML_TEXTBOX):
            for t_elem in txbx.iter(_W_T):
                t = (t_elem.text or "").strip()
                if t:
                    texts.append(t)
        return texts

    # ------------------------------------------------------------------
    # Footnotes and endnotes
    # ------------------------------------------------------------------

    def _extract_footnotes_endnotes(self, document: Document) -> List[str]:
        """Extract text from footnotes and endnotes if present."""
        from lxml import etree  # lxml is a python-docx transitive dep

        _FOOTNOTES_REL = (
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
        )
        _ENDNOTES_REL = (
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"
        )
        parts_to_check = [
            (_FOOTNOTES_REL, _W_FOOTNOTE),
            (_ENDNOTES_REL, _W_ENDNOTE),
        ]
        texts: List[str] = []
        for rel_type, note_tag in parts_to_check:
            try:
                part = document.part.part_related_by(rel_type)
                root = etree.fromstring(part.blob)
                for note in root.iter(note_tag):
                    # Skip built-in separator notes (id <= 0).
                    note_id = int(note.get(_W_ID, "1"))
                    if note_id <= 0:
                        continue
                    t = "".join(
                        (elem.text or "") for elem in note.iter(_W_T)
                    ).strip()
                    if t:
                        texts.append(t)
            except (KeyError, AttributeError, TypeError):
                pass  # Part doesn't exist in this document.
        return texts

    # ------------------------------------------------------------------
    # Legacy .doc support
    # ------------------------------------------------------------------

    def _parse_doc_file(self, file_path: str) -> str:
        """Attempt to parse a legacy .doc file via python-docx (limited support)."""
        logger.warning("Legacy .doc format detected: %s", file_path)
        try:
            document = Document(file_path)
            texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            if texts:
                logger.info("Parsed .doc file via docx: %s", file_path)
                return "\n".join(texts)
        except Exception as e:
            raise FileParsingError(
                "Error parsing Word document.", original_exception=e
            ) from e

        raise FileParsingError(
            "Cannot parse legacy .doc file. "
            "Convert to .docx format first (e.g. with LibreOffice)."
        )
